from PyQt6.QtWidgets import *
from PyQt6.QtGui import *
from PyQt6.QtCore import *
from PyQt6.QtWebEngineWidgets import *
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
from docx import Document
from docx.shared import Pt, RGBColor
import os
import re
import webbrowser
import sys
import threading
import requests
from flask import Flask, request, session
from authlib.integrations.flask_client import OAuth
from authlib.integrations.requests_client import OAuth2Session
from requests_oauthlib import OAuth2Session
import threading
from auth import app, github
from urllib.parse import quote
import codecs
from base64 import b64decode


def get_username_from_about_file():
    with open("about.txt", "r") as file:
        lines = file.readlines()
        for line in lines:
            if line.startswith("Username:"):
                return line.strip().split(":")[1].strip()
    return None


class TextEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Text Editor")
        self.setGeometry(100, 100, 800, 600)

        self.tab_widget = QTabWidget(self)
        self.setCentralWidget(self.tab_widget)

        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self.close_tab)

        self.init_menu()
        self.init_toolbar()
        self.init_tab_bar()
        self.open_empty_tab()
        self.text_area = QTextEdit()

        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)

        # Statistik-Label erstellen
        self.stats_label = QLabel("Word count: 0 | Character count: 0", self)
        self.status_bar.addWidget(self.stats_label)
        self.line_number_label = QLabel()
        self.statusBar().addPermanentWidget(self.line_number_label)


    def init_menu(self):
        menubar = self.menuBar()
        # Infos-Menü
        info_menu = menubar.addMenu('Infos')
        developer_action = QAction('Developer', self)
        developer_action.triggered.connect(self.show_developer_action)
        info_menu.addAction(developer_action)

        about_action = QAction('About', self)
        about_action.triggered.connect(self.show_info_dock)
        info_menu.addAction(about_action)

        # QDockWidget für Infos erstellen
        self.info_dock = QDockWidget("Infos", self)
        self.info_dock.setAllowedAreas(Qt.DockWidgetArea.BottomDockWidgetArea)
        self.addDockWidget(Qt.DockWidgetArea.BottomDockWidgetArea, self.info_dock)

        # Layout für den Inhalt des QDockWidget erstellen
        dock_content = QWidget()
        dock_layout = QVBoxLayout(dock_content)
        self.label_word_count = QLabel()
        self.label_char_count = QLabel()
        dock_layout.addWidget(self.label_word_count)
        dock_layout.addWidget(self.label_char_count)

        # QDockWidget mit dem Layout verbinden
        self.info_dock.setWidget(dock_content)

        # Das QDockWidget initial verstecken
        self.info_dock.hide()


        file_menu = menubar.addMenu("File")
        open_action = QAction("Open", self)
        open_action.triggered.connect(self.open_file)
        file_menu.addAction(open_action)

        new_tab_action = QAction("New Tab", self)
        new_tab_action.setShortcut(QKeySequence("Ctrl+T")) 
        new_tab_action.triggered.connect(self.open_new_tab)
        file_menu.addAction(new_tab_action)

        search_action = QAction("Search Word", self)
        search_action.setShortcut(QKeySequence("Ctrl+F"))
        search_action.triggered.connect(self.show_search_dialog)
        file_menu.addAction(search_action)

        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(QApplication.quit)
        file_menu.addAction(exit_action)

        save_menu = menubar.addMenu("Save")
        save_action = QAction("Save", self)
        save_action.setShortcut(QKeySequence("Ctrl+S"))
        save_action.triggered.connect(self.save_file)
        save_menu.addAction(save_action)

        export_docx_action = QAction("Export as DOCX", self)
        export_docx_action.triggered.connect(self.export_as_docx)
        save_menu.addAction(export_docx_action)

        export_txt_action = QAction("Export as TXT", self)
        export_txt_action.triggered.connect(self.export_as_txt)
        save_menu.addAction(export_txt_action)

        print_action = QAction("Print", self)
        print_action.setShortcut(QKeySequence("Ctrl+P"))
        print_action.triggered.connect(self.print_document)
        save_menu.addAction(print_action)

        edit_menu = menubar.addMenu("Edit")
        undo_action = QAction("Undo", self)
        undo_action.setShortcut(QKeySequence.StandardKey.Undo)
        undo_action.triggered.connect(self.undo)
        edit_menu.addAction(undo_action)

        redo_action = QAction("Redo", self)
        redo_action.setShortcut(QKeySequence.StandardKey.Redo)
        redo_action.triggered.connect(self.redo)
        edit_menu.addAction(redo_action)

        # Move "Login" action to the "Projects" menu
        projects_menu = menubar.addMenu("Projects")
        login_action = QAction("Login", self)
        login_action.triggered.connect(self.start_webserver)
        projects_menu.addAction(login_action)

        projects = self.load_projects()
        if projects:
            # Add a separator between "Login" and the projects
            projects_menu.addSeparator()

        self.set_style_options()
            
        # Add projects to the menu
        for project in projects:
            project_action = QAction(project, self)
            project_action.triggered.connect(lambda _, p=project: self.open_project(p))
            projects_menu.addAction(project_action)
    def update_status_bar(self):
        current_widget = self.tab_widget.currentWidget()
        if isinstance(current_widget, QTextEdit):
            content = current_widget.toPlainText()

            # Word und Zeichen zählen (Apostrophe als Teil eines Worts berücksichtigen)
            word_count = len(re.findall(r'\b\w+\'?\w*\b', content))
            char_count = len(content)

            # Zeilennummer abrufen
            cursor = current_widget.textCursor()
            line_number = cursor.blockNumber() + 1  # BlockNumber() gibt 0-basierte Nummer zurück, daher +1

            # Statusleisteninformationen aktualisieren
            self.stats_label.setText(f"Word count: {word_count} | Character count: {char_count}")

            # Zeilennummer ganz rechts in der Statusleiste anzeigen
            self.line_number_label.setText(f"Line: {line_number}")
        else:
            self.stats_label.setText("")
            self.line_number_label.setText("")

    def show_developer_action(self):
        developer_info_dialog = QDialog(self)
        developer_info_dialog.setWindowTitle("Entwickler")

        # GitHub-Repository-Informationen
        repo_owner = "SchBenedikt"
        repo_name = "Text-Editor"
        repo_url = f"https://api.github.com/repos/{repo_owner}/{repo_name}/contributors"

        # Anfrage an die GitHub-API senden
        response = requests.get(repo_url)
        if response.status_code == 200:
            contributors = response.json()

            # QTableWidget für die Entwicklerinformationen erstellen
            table = QTableWidget()
            table.setRowCount(len(contributors))
            table.setColumnCount(1)  # Eine Spalte für den Benutzernamen
            table.setHorizontalHeaderLabels(["Username"])

            # Tabelle mit Daten füllen
            for row, contributor in enumerate(contributors):
                username = contributor.get("login", "Unknown")

                # Daten in die Tabelle einfügen
                table.setItem(row, 0, QTableWidgetItem(username))

            # Tabelle zum Layout des Dialogs hinzufügen
            layout = QVBoxLayout(developer_info_dialog)
            layout.addWidget(table)

            # Tabelle an die Größe des Inhalts anpassen
            table.resizeColumnsToContents()

            # Die Bearbeitung der Zellen in der Tabelle deaktivieren
            table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)

            # Tabelle soll die gesamte Breite des Dialogs einnehmen
            header = table.horizontalHeader()
            header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)

            # Das Dialogfenster in der Mitte des Hauptfensters positionieren
            rect = self.geometry()
            developer_info_dialog.move(rect.center() - developer_info_dialog.rect().center())

            # Die Modalität des Fensters auf Anwendungsmodalität setzen
            developer_info_dialog.setWindowModality(Qt.WindowModality.ApplicationModal)

            # Das QDialog anzeigen
            developer_info_dialog.exec()
    def show_info_dock(self):
        # Replace with your actual repository and author
        repo_owner = "SchBenedikt"
        repo_name = "Text-Editor"

        # Fetch the latest release information from GitHub
        release_url = f"https://api.github.com/repos/{repo_owner}/{repo_name}/releases/latest"
        response = requests.get(release_url)
        if response.status_code == 200:
            latest_version = response.json().get("tag_name", "vYYYY.MM.DD")  # Replace with a default version
        else:
            latest_version = "vYYYY.MM.DD"  # Replace with a default version

        # Compare the latest version with the current version
        current_version = "v2024.01.04"  # Replace with your actual version

        # Create a new QDialog for displaying version information
        info_dialog = QDialog(self)
        info_dialog.setWindowTitle("About")

        # Create QLabel widgets for displaying version information
        if latest_version != current_version:
            label_version = QLabel(f"Current Version: {current_version}\n\nNew Version: {latest_version}", info_dialog)
            button_open_release = QPushButton("New release available", info_dialog)
            button_open_release.clicked.connect(lambda: QDesktopServices.openUrl(QUrl(release_url)))
        else:
            label_version = QLabel(f"{current_version}", info_dialog)
            button_open_release = None

        # Add widgets to layout
        layout = QVBoxLayout(info_dialog)
        layout.addWidget(label_version)
        if button_open_release:
            layout.addWidget(button_open_release)

        # Adjust the size of the window based on its contents plus 10 pixels
        info_dialog.adjustSize()
        info_dialog.setFixedSize(info_dialog.width() + 10, info_dialog.height() + 10)

        # Center the dialog on the main window
        rect = self.geometry()
        info_dialog.move(rect.center() - info_dialog.rect().center())

        # Set the window modality to be application modal
        info_dialog.setWindowModality(Qt.WindowModality.ApplicationModal)

        # Show the QDialog
        info_dialog.exec()
    def open_project(self, project):
        username = get_username_from_about_file()
        if username:
            repo_url = f"https://api.github.com/repos/{username}/{project}/contents"
            try:
                response = requests.get(repo_url)
                response.raise_for_status()  # Check if the request was successful

                files = [file_info["name"] for file_info in response.json()]

                # Display the file names in a QMessageBox
                if files:
                    selected_file, ok = QInputDialog.getItem(self, "Select File", "Select a file to open:", files, 0,
                                                             False)
                    if ok and selected_file:
                        # Fetch the content of the selected file
                        file_content_url = f"https://api.github.com/repos/{username}/{project}/contents/{quote(selected_file)}"
                        content_response = requests.get(file_content_url)
                        content_response.raise_for_status()

                        # Decode the base64-encoded content
                        content = b64decode(content_response.json()["content"]).decode("utf-8")

                        current_widget = self.tab_widget.currentWidget()
                        current_widget.setPlainText(content)
                        self.set_tab_title(current_widget, selected_file)
                else:
                    QMessageBox.warning(self, "No Files", f"There are no files in {project}.")
            except requests.RequestException as e:
                QMessageBox.warning(self, "Error", f"Error fetching project files: {str(e)}")
        else:
            # Handle the case when username is not available
            pass

    def load_projects(self):
        projects = []
        with open("projects.txt", "r") as file:
            for line in file:
                project = line.strip()
                if project:
                    projects.append(project)
        return projects

    def start_webserver(self):
        def run_flask_app():
            app.run(host="127.0.0.1", port=5000)

        flask_thread = threading.Thread(target=run_flask_app)
        flask_thread.start()

        # Open web browser to localhost:5000
        url = "http://127.0.0.1:5000"
        webbrowser.open(url)
    def undo(self):
        current_widget = self.tab_widget.currentWidget()
        current_widget.undo()

    def redo(self):
        current_widget = self.tab_widget.currentWidget()
        current_widget.redo()

    def init_toolbar(self):
        toolbar = QToolBar(self)
        self.addToolBar(toolbar)

        toolbar.setIconSize(QSize(20, 20))

        bold_action = QAction(QIcon("images/bold.png"), "Bold", self)
        bold_action.triggered.connect(self.bold_text)
        toolbar.addAction(bold_action)

        italic_action = QAction(QIcon("images/italic.png"), "Italic", self)
        italic_action.triggered.connect(self.italic_text)
        toolbar.addAction(italic_action)

        underline_action = QAction(QIcon("images/underline.png"), "Underline", self)
        underline_action.triggered.connect(self.underline_text)
        toolbar.addAction(underline_action)

        increase_font_action = QAction(QIcon("images/increase_font.png"), "Increase Font Size", self)
        increase_font_action.triggered.connect(self.increase_font_size)
        toolbar.addAction(increase_font_action)

        decrease_font_action = QAction(QIcon("images/decrease_font.png"), "Decrease Font Size", self)
        decrease_font_action.triggered.connect(self.decrease_font_size)
        toolbar.addAction(decrease_font_action)

        font_combobox = QFontComboBox(self)
        font_combobox.setCurrentFont(QFont("TkDefaultFont"))
        font_combobox.currentFontChanged.connect(self.change_font)
        toolbar.addWidget(font_combobox)

        change_color_action = QAction(QIcon("images/change_color.png"), "Change Text Color", self)
        change_color_action.triggered.connect(self.change_text_color)
        toolbar.addAction(change_color_action)

        set_text_background_color = QAction(QIcon("images/change_bg_color.png"), "Change Background Color", self)
        set_text_background_color.triggered.connect(self.set_text_background_color)
        toolbar.addAction(set_text_background_color)

       
    def init_tab_bar(self):
        add_tab_button = QToolButton(self)
        add_tab_button.setText("+")
        add_tab_button.setStyleSheet("QToolButton { font-size: 20px; }")
        add_tab_button.clicked.connect(self.open_new_tab)
        self.tab_widget.setCornerWidget(add_tab_button)

    def open_file(self):
        file, _ = QFileDialog.getOpenFileName(self, "Open File")
        if file:
            encodings = ["utf-8", "latin-1", "cp1252"]
            content = None
            for encoding in encodings:
                try:
                    with open(file, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            if content is not None:
                current_widget = self.tab_widget.currentWidget()
                current_widget.setPlainText(content)
                self.set_tab_title(current_widget, file)
            else:
                QMessageBox.warning(self, "Open File", "Unable to open the file.")


    def save_file(self):
        # Show a dialog to choose the save option
        options = ['Save locally', 'Save on GitHub']
        choice, ok = QInputDialog.getItem(self, 'Save Option', 'Choose a save option:', options, 0, False)

        if ok:
            current_widget = self.tab_widget.currentWidget()
            content = current_widget.toPlainText()

            if choice == 'Save locally':
                self.save_locally(content)
            elif choice == 'Save on GitHub':
                if os.path.exists('upload_data.txt'):
                    github_username, access_token = self.read_upload_data()
                    repo_name = self.get_user_repository(github_username, access_token)
                    if repo_name:
                        self.save_to_github(content, github_username, access_token, repo_name)
                    else:
                        QMessageBox.warning(self, 'GitHub Save', 'No repositories found for the given GitHub username.')
                else:
                    self.ask_github_credentials_and_save(content)


    def save_locally(self, content):
        file_dialog = QFileDialog()
        file_name, _ = file_dialog.getSaveFileName(self, "Save File Locally", "", "All Files (*)")

        if file_name:
            with open(file_name, "w") as f:
                f.write(content)
            QMessageBox.information(self, "Save Successful", "File saved locally successfully.")

    def ask_github_credentials_and_save(self, content):
        github_username, ok1 = QInputDialog.getText(self, 'GitHub Credentials', 'Enter your GitHub username:')
        access_token, ok2 = QInputDialog.getText(self, 'GitHub Credentials', 'Enter your GitHub personal access token:')

        if ok1 and ok2:
            repositories = self.get_user_repositories(github_username, access_token)

            if repositories:
                repository_name, ok3 = QInputDialog.getItem(self, 'Select Repository', 'Choose a GitHub repository:', repositories, 0, False)

                if ok3:
                    with open('upload_data.txt', 'w') as file:
                        file.write(f"{github_username}\n{access_token}\n{repository_name}")

                    self.save_to_github(content, github_username, access_token, repository_name)
                else:
                    QMessageBox.warning(self, 'GitHub Save', 'No GitHub repository selected. Unable to save on GitHub.')
            else:
                QMessageBox.warning(self, 'GitHub Save', 'No repositories found for the given GitHub username.')
        else:
            QMessageBox.warning(self, 'GitHub Save', 'GitHub credentials not provided. Unable to save on GitHub.')
    def get_user_repositories(self, github_username, access_token):
        api_url = f"https://api.github.com/users/{github_username}/repos"
        headers = {"Authorization": f"token {access_token}"}

        try:
            response = requests.get(api_url, headers=headers)
            repositories = [repo['name'] for repo in response.json()]
            return repositories
        except requests.RequestException as e:
            print(f"Error getting GitHub repositories: {e}")
            return None


    def save_to_github(self, content, github_username, access_token, repo_name):
        # Show a dialog to input the desired file name for GitHub
        custom_github_filename, ok = QInputDialog.getText(self, 'GitHub File Name', 'Enter the desired file name for GitHub (with extension):')

        if ok and custom_github_filename:
            api_url = f"https://api.github.com/repos/{github_username}/{repo_name}/contents/{custom_github_filename}"
            headers = {"Authorization": f"token {access_token}"}

            data = {
                "message": "Upload file via schBenedikt's Text Editor",
                "content": codecs.encode(content.encode("utf-8"), "base64").decode("utf-8"),
                "sha": self.get_sha_from_github(custom_github_filename, github_username, access_token, repo_name)
            }

            response = requests.put(api_url, headers=headers, json=data)

            if response.status_code == 200:
                print(f"File '{custom_github_filename}' uploaded to GitHub successfully.")
            else:
                print(f"Unable to upload file to GitHub. Status code: {response.status_code}, Message: {response.text}")

    def get_user_repositories(self, github_username, access_token):
        api_url = f"https://api.github.com/users/{github_username}/repos"
        headers = {"Authorization": f"token {access_token}"}

        try:
            response = requests.get(api_url, headers=headers)
            repositories = [repo['name'] for repo in response.json()]
            return repositories
        except requests.RequestException as e:
            print(f"Error getting GitHub repositories: {e}")
            return None

    def get_user_repository(self, github_username, access_token):
        repositories = self.get_user_repositories(github_username, access_token)
        if repositories:
            repository_name, ok = QInputDialog.getItem(self, 'Select Repository', 'Choose a GitHub repository:', repositories, 0, False)
            if ok:
                return repository_name
        return None


    def get_filename_from_path(self, file_path):
        return os.path.basename(file_path)

    def read_upload_data(self):
        with open('upload_data.txt', 'r') as file:
            lines = file.readlines()
            github_username = lines[0].strip()
            access_token = lines[1].strip()
            return github_username, access_token


    def upload_to_github(self, content, github_filename, github_username, access_token, repo_name):
        # Extrahiere nur den Dateinamen aus dem vollständigen Pfad
        github_filename = os.path.basename(github_filename)

        api_url = f"https://api.github.com/repos/{github_username}/{repo_name}/contents/{github_filename}"
        headers = {"Authorization": f"token {access_token}"}

        data = {
            "message": "Update file via script",
            "content": codecs.encode(content.encode("utf-8"), "base64").decode("utf-8"),
            "sha": self.get_sha_from_github(github_filename, github_username, access_token, repo_name)
        }

        response = requests.put(api_url, headers=headers, json=data)

        if response.status_code == 200:
            print(f"File '{github_filename}' uploaded to GitHub successfully.")
        else:
            print(f"Unable to upload file to GitHub. Status code: {response.status_code}, Message: {response.text}")

    def get_sha_from_github(self, github_filename, github_username, access_token, repo_name):
        api_url = f"https://api.github.com/repos/{github_username}/{repo_name}/contents/{github_filename}"
        headers = {"Authorization": f"token {access_token}"}

        response = requests.get(api_url, headers=headers)

        if response.status_code == 200:
            sha = response.json().get("sha")
            return sha
        else:
            print(f"Unable to get SHA from GitHub. Status code: {response.status_code}, Message: {response.text}")
            return None

    def load_github_credentials(self):
        # Read GitHub credentials from upload_data.txt
        try:
            with open('upload_data.txt', 'r') as file:
                lines = file.readlines()
                github_username = lines[0].strip()
                access_token = lines[1].strip()
                repo_name = lines[2].strip()
                return github_username, access_token, repo_name
        except FileNotFoundError:
            print("upload_data.txt not found.")
            return None, None, None

    def set_tab_title(self, text_widget, file_path):
        # Set the tab title to the file name
        file_name = QFileInfo(file_path).fileName()
        index = self.tab_widget.indexOf(text_widget)
        self.tab_widget.setTabText(index, file_name)
        # Store the file path in the text_widget for later use
        setattr(text_widget, "file_path", file_path)

    def export_as_docx(self):
        file, _ = QFileDialog.getSaveFileName(self, "Export as DOCX", filter="*.docx")
        if file:
            doc = Document()
            current_widget = self.tab_widget.currentWidget()
            # content = current_widget.toPlainText()
            paragraph = doc.add_paragraph()
            runs = self.get_runs_with_formatting(current_widget)
            for run_text, font_format in runs:
                run = paragraph.add_run(run_text)
                self.apply_formatting(run, font_format)
            doc.save(file)
            QMessageBox.information(self, "Export as DOCX", "File exported successfully.")

    def export_as_txt(self):
        file, _ = QFileDialog.getSaveFileName(self, "Export as TXT", filter="*.txt")
        if file:
            current_widget = self.tab_widget.currentWidget()
            content = current_widget.toPlainText()
            try:
                with open(file, "w") as f:
                    f.write(content)
                QMessageBox.information(self, "Export as TXT", "File exported successfully.")
            except:
                QMessageBox.warning(self, "Export as TXT", "Unable to export the file.")

    def apply_formatting(self, run, font_format):
        font = run.font
        if font_format["bold"]:
            font.bold = True
        if font_format["italic"]:
            font.italic = True
        if font_format["underline"]:
            font.underline = True
        if font_format["color"]:
            rgb_color = QColor(font_format["color"])
            # font.color.rgb = rgb_color.rgb()
            font.color.rgb = RGBColor(rgb_color.red(), rgb_color.green(), rgb_color.blue())
        font.size = Pt(15)

    def get_runs_with_formatting(self, text_widget):
        # cursor = QTextCursor(self.text_area.document())
        cursor = QTextCursor(text_widget.document())
        cursor.setPosition(0)
        cursor.movePosition(QTextCursor.End, QTextCursor.KeepAnchor)
        selected_text = cursor.selection().toPlainText()

        runs = []
        start = 0
        for char in selected_text:
            cursor.setPosition(start)
            cursor.movePosition(QTextCursor.NextCharacter, QTextCursor.KeepAnchor)
            char_format = cursor.charFormat()
            runs.append((char, {
                "bold": char_format.font().bold(),
                "italic": char_format.font().italic(),
                "underline": char_format.font().underline(),
                "color": char_format.foreground().color().name()
            }))
            start += 1
        return runs


    def is_unsaved_changes(self, text_widget):
        if isinstance(text_widget, QTextEdit):  # Check if text_widget is a QTextEdit instance
            content = text_widget.toPlainText()
            return content != "" and content != self.get_file_content(text_widget)
        return False
    def get_file_content(self, text_widget):
        file_path = getattr(text_widget, "file_path", None)
        if file_path:
            try:
                with open(file_path, "r") as file:
                    content = file.read()
                    return content
            except FileNotFoundError:
                pass
        return ""

    def bold_text(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font.setBold(not font.bold())
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def italic_text(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font.setItalic(not font.italic())
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def underline_text(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font.setUnderline(not font.underline())
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def increase_font_size(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font_size = font.pointSize()
        font_size += 1
        font.setPointSize(font_size)
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def decrease_font_size(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font_size = font.pointSize()
        font_size -= 1
        font.setPointSize(font_size)
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def change_font(self, font):
        current_widget = self.tab_widget.currentWidget()
        text_cursor = current_widget.textCursor()
        format = text_cursor.charFormat()

        font_name = font.family()
        format.setFontFamily(font_name)
        text_cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def change_text_color(self):
        print("Before opening color dialog")
        color = QColorDialog.getColor(parent=self)
        print("After opening color dialog")

        if color.isValid():
            current_widget = self.tab_widget.currentWidget()
            cursor = current_widget.textCursor()
            format = cursor.charFormat()
            print(f"Cursor: {cursor}")
            print(f"Format: {format}")

            print(f"Color: {color}")
            format.setForeground(color)
            print(f"Foreground color: {format.foreground()}")

            cursor.mergeCharFormat(format)
            current_widget.setFocus()
        elif not color.isValid():
            print("Invalid color")


    def set_text_background_color(self, color):
        color = QColorDialog.getColor(parent=self)
        if color.isValid():
            current_widget = self.tab_widget.currentWidget()
            cursor = current_widget.textCursor()
            format = cursor.charFormat()
            format.setBackground(color)
            cursor.mergeCharFormat(format)
            current_widget.setFocus()

    def open_new_tab(self):
        options = ["New File", "Open File", "Chat"]
        selected_option, ok = QInputDialog.getItem(self, "New File, Open File, or Chat?", "Choose an option:", options, 0, False)

        if ok:
            if selected_option == "New File":
                dialog = QFileDialog(self)
                dialog.setFileMode(QFileDialog.AnyFile)

                options = QFileDialog.Options()
                options |= QFileDialog.DontUseNativeDialog
                fileName, _ = dialog.getSaveFileName(self, "New File", "", "All Files (*);;Text Files (*.txt);;Python Files (*.py)", options=options)

                if fileName:
                    if fileName.endswith(".txt"):
                        self.open_text_file_in_tab(fileName)
                    elif fileName.endswith(".py"):
                        self.open_python_file_in_tab(fileName)
                    else:
                        self.open_generic_file_in_tab(fileName)
                else:
                    self.open_empty_tab()
            elif selected_option == "Open File":
                dialog = QFileDialog(self)
                dialog.setFileMode(QFileDialog.ExistingFile)

                options = QFileDialog.Options()
                options |= QFileDialog.DontUseNativeDialog
                fileName, _ = dialog.getOpenFileName(self, "Open File", "", "All Files (*);;Text Files (*.txt);;Python Files (*.py)", options=options)

                if fileName:
                    if fileName.endswith(".txt"):
                        self.open_text_file_in_tab(fileName)
                    elif fileName.endswith(".py"):
                        self.open_python_file_in_tab(fileName)
                    else:
                        self.open_generic_file_in_tab(fileName)
            elif selected_option == "Chat":
                self.open_chat_tab()

    def open_chat_tab(self):
        chat_view = QWebEngineView()
        chat_view.setUrl(QUrl("https://platform.openai.com/"))
        chat_widget = QWidget()
        layout = QVBoxLayout(chat_widget)
        layout.addWidget(chat_view)

        self.tab_widget.addTab(chat_widget, "Chat")


    def open_text_file_in_tab(self, file_path):
        with open(file_path, "r") as file:
            content = file.read()

        text_area = QTextEdit()
        text_area.setFont(QFont("TkDefaultFont"))
        text_area.textChanged.connect(self.update_tab_title)
        text_area.setPlainText(content)
        self.tab_widget.addTab(text_area, os.path.basename(file_path))
        self.tab_widget.setCurrentWidget(text_area)

    def open_python_file_in_tab(self, file_path):
        # Implementiere entsprechende Logik, um Python-Dateien zu öffnen
        pass

    def open_generic_file_in_tab(self, file_path):
        # Implementiere entsprechende Logik, um generische Dateien zu öffnen
        pass
    def open_empty_tab(self):
        text_area = QTextEdit()
        text_area.setFont(QFont("TkDefaultFont"))
        text_area.textChanged.connect(self.update_tab_title)
        text_area.textChanged.connect(self.update_status_bar)
        self.tab_widget.addTab(text_area, "Untitled")
        self.tab_widget.setCurrentWidget(text_area)

    def open_new_empty_tab(self):
        text_area = QTextEdit()
        text_area.setFont(QFont("TkDefaultFont"))
        text_area.textChanged.connect(self.update_tab_title)
        text_area.textChanged.connect(self.update_status_bar)
        self.tab_widget.addTab(text_area, "Untitled")
        self.tab_widget.setCurrentWidget(text_area)
    def close_tab(self, index):
        text_area = self.tab_widget.widget(index)

        if self.is_unsaved_changes(text_area):
            reply = QMessageBox.question(self, "Unsaved Changes",
                                         "There are unsaved changes. Do you want to save before closing the tab?",
                                         QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel)
            if reply == QMessageBox.Save:
                self.save_file()
            elif reply == QMessageBox.Cancel:
                return

        self.tab_widget.removeTab(index)
    def closeEvent(self, event):
        current_widget = self.tab_widget.currentWidget()
        if self.is_unsaved_changes(current_widget):
            reply = QMessageBox.question(self, "Unsaved Changes",
                                         "There are unsaved changes. Do you want to save before exiting?",
                                         QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel)
            if reply == QMessageBox.Save:
                self.save_file()
            elif reply == QMessageBox.Cancel:
                event.ignore()
                return
        event.accept()

        # Check if it's the last tab being closed
        if self.tab_widget.count() == 0:
            # Close the entire application
            self.close()
    def update_tab_title(self):
        current_widget = self.tab_widget.currentWidget()
        current_index = self.tab_widget.currentIndex()
        if self.is_unsaved_changes(current_widget):
            file_path = getattr(current_widget, "file_path", None)
            if file_path:
                file_name = os.path.basename(file_path)
                self.tab_widget.setTabText(current_index, file_name + " *")
            else:
                self.tab_widget.setTabText(current_index, "Unsaved *")
        else:
            file_path = getattr(current_widget, "file_path", None)
            if file_path:
                file_name = os.path.basename(file_path)
                self.tab_widget.setTabText(current_index, file_name)
            else:
                self.tab_widget.setTabText(current_index, "Untitled")
    def print_document(self):
        printer = QPrinter(QPrinter.HighResolution)
        printer.setPageSize(QPrinter.PageSize.A4)

        dialog = QPrintDialog(printer, self)
        if dialog.exec() == QPrintDialog.Accepted:
            current_widget = self.tab_widget.currentWidget()
            current_widget.print(printer)
    def search_word(self, word):
        current_widget = self.tab_widget.currentWidget()
        text_edit = current_widget
        cursor = QTextCursor(text_edit.document())

        if cursor.hasSelection():
            cursor.clearSelection()

        found = False
        while True:
            cursor = text_edit.document().find(word, cursor)

            if cursor.isNull():
                break

            found = True

            cursor.select(QTextCursor.WordUnderCursor)
            text_edit.setTextCursor(cursor)
            text_edit.ensureCursorVisible()

            reply = QMessageBox.question(self, "Word Found",
                                         f"The word '{word}' was found in the document. Do you want to continue searching?",
                                         QMessageBox.Yes | QMessageBox.No)
            if reply == QMessageBox.No:
                break

        if not found:
            QMessageBox.information(self, "Word Not Found", f"The word '{word}' was not found in the document.")

    def show_search_dialog(self):
        word, ok = QInputDialog.getText(self, "Search Word", "Enter the word to search:")
        if ok:
            self.search_word(word)

    def set_style_options(self):
        style_sheet = """
QTabWidget::pane {
    background-color: #FFFFFF;
}

QTabWidget::tab-bar {
    alignment: left;
    height: auto;
}

QTabBar::tab {
    background-color: #FFFFFF;
    color: #050000;
    border: 1px solid #C0C0C0;
    padding: 10px 0;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    width: 100px;
    height: 10px;
    text-align: center; 
    transition: background-color 0.3s ease, color 0.3s ease;
}

QTabBar::tab:selected {
    background-color: #FFFFFF;
    color: #000000;
    border-bottom-color: #FFFFFF;
}

QTabBar::tab:!selected {
    margin-top: 2px;
    background-color: #ebfcfc;
}

QTabBar::tab:hover {
    background-color: #a5faf8; 
    color: #303030; 
    cursor: pointer; 
}


        """
        self.setStyleSheet(style_sheet)


if __name__ == "__main__":
    app_thread = threading.Thread(target=app.run, kwargs={"host": "localhost", "port": 5000})
    app_thread.daemon = True
    app_thread.start()

    app_pyqt = QApplication(sys.argv)
    window = TextEditor()
    window.show()

    sys.exit(app_pyqt.exec())
