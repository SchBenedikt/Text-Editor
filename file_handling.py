import os
from PyQt6.QtWidgets import QFileDialog, QMessageBox
from PyQt6.QtPrintSupport import QPrinter, QPrintDialog
from docx import Document
from docx.shared import Pt, RGBColor
import codecs

def open_file(text_widget):
    file, _ = QFileDialog.getOpenFileName(text_widget, "Open File")
    if file:
        encodings = ["utf-8", "latin-1", "cp1252"]
        content = None
        for encoding in encodings:
            try:
                with open(file, "r", encoding=encoding) as f:
                    content = f.read()
                break
            except UnicodeDecodeError:
                continue
        if content is not None:
            text_widget.setPlainText(content)
            return file
        else:
            QMessageBox.warning(text_widget, "Open File", "Unable to open the file.")
    return None

def save_file(text_widget, content, save_locally, save_to_github):
    options = ['Save locally', 'Save on GitHub']
    choice, ok = QFileDialog.getItem(text_widget, 'Save Option', 'Choose a save option:', options, 0, False)

    if ok:
        if choice == 'Save locally':
            save_locally(content)
        elif choice == 'Save on GitHub':
            save_to_github(content)

def save_locally(content):
    file_dialog = QFileDialog()
    file_name, _ = file_dialog.getSaveFileName(None, "Save File Locally", "", "All Files (*)")

    if file_name:
        with open(file_name, "w") as f:
            f.write(content)
        QMessageBox.information(None, "Save Successful", "File saved locally successfully.")

def export_as_docx(text_widget):
    file, _ = QFileDialog.getSaveFileName(text_widget, "Export as DOCX", filter="*.docx")
    if file:
        doc = Document()
        paragraph = doc.add_paragraph()
        runs = get_runs_with_formatting(text_widget)
        for run_text, font_format in runs:
            run = paragraph.add_run(run_text)
            apply_formatting(run, font_format)
        doc.save(file)
        QMessageBox.information(text_widget, "Export as DOCX", "File exported successfully.")

def export_as_txt(text_widget):
    file, _ = QFileDialog.getSaveFileName(text_widget, "Export as TXT", filter="*.txt")
    if file:
        content = text_widget.toPlainText()
        try:
            with open(file, "w") as f:
                f.write(content)
            QMessageBox.information(text_widget, "Export as TXT", "File exported successfully.")
        except:
            QMessageBox.warning(text_widget, "Export as TXT", "Unable to export the file.")

def apply_formatting(run, font_format):
    font = run.font
    if font_format["bold"]:
        font.bold = True
    if font_format["italic"]:
        font.italic = True
    if font_format["underline"]:
        font.underline = True
    if font_format["color"]:
        rgb_color = QColor(font_format["color"])
        font.color.rgb = RGBColor(rgb_color.red(), rgb_color.green(), rgb_color.blue())
    font.size = Pt(15)

def get_runs_with_formatting(text_widget):
    cursor = QTextCursor(text_widget.document())
    cursor.setPosition(0)
    cursor.movePosition(QTextCursor.MoveOperation.End, QTextCursor.MoveMode.KeepAnchor)
    selected_text = cursor.selection().toPlainText()

    runs = []
    start = 0
    for char in selected_text:
        cursor.setPosition(start)
        cursor.movePosition(QTextCursor.MoveOperation.NextCharacter, QTextCursor.MoveMode.KeepAnchor)
        char_format = cursor.charFormat()
        runs.append((char, {
            "bold": char_format.font().bold(),
            "italic": char_format.font().italic(),
            "underline": char_format.font().underline(),
            "color": char_format.foreground().color().name()
        }))
        start += 1
    return runs

def print_document(text_widget):
    printer = QPrinter(QPrinter.HighResolution)
    printer.setPageSize(QPrinter.PageSize.A4)

    dialog = QPrintDialog(printer, text_widget)
    if dialog.exec() == QPrintDialog.Accepted:
        text_widget.print(printer)
