from PyQt5.QtWidgets import *
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from docx import Document
from docx.shared import Pt
import os

class TextEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Text Editor")
        self.setGeometry(100, 100, 800, 600)

        self.tab_widget = QTabWidget(self)
        self.setCentralWidget(self.tab_widget)

        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self.close_tab)

        self.init_menu()
        self.init_toolbar()
        self.init_tab_bar()
        # QApplication.setStyle(QStyleFactory.create("Fusion"))

        # Set additional Fusion style options for Windows 10/11-like appearance
        self.set_style_options()
        self.open_new_tab()

    def init_menu(self):
        menubar = self.menuBar()

        file_menu = menubar.addMenu("File")

        open_action = QAction("Open", self)
        open_action.triggered.connect(self.open_file)
        file_menu.addAction(open_action)

        new_tab_action = QAction("New Tab", self)
        new_tab_action.triggered.connect(self.open_new_tab)
        file_menu.addAction(new_tab_action)

        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(qApp.quit)
        file_menu.addAction(exit_action)

        save_menu = menubar.addMenu("Save")
        save_action = QAction("Save", self)
        save_action.triggered.connect(self.save_file)
        save_menu.addAction(save_action)

        export_docx_action = QAction("Export as DOCX", self)
        export_docx_action.triggered.connect(self.export_as_docx)
        save_menu.addAction(export_docx_action)

        export_txt_action = QAction("Export as TXT", self)
        export_txt_action.triggered.connect(self.export_as_txt)
        save_menu.addAction(export_txt_action)

    def init_toolbar(self):
        toolbar = QToolBar(self)
        self.addToolBar(toolbar)

        toolbar.setIconSize(QSize(20, 20)) 

        bold_action = QAction(QIcon("bold.png"), "Bold", self)
        bold_action.triggered.connect(self.bold_text)
        toolbar.addAction(bold_action)

        italic_action = QAction(QIcon("italic.png"), "Italic", self)
        italic_action.triggered.connect(self.italic_text)
        toolbar.addAction(italic_action)

        underline_action = QAction(QIcon("underline.png"), "Underline", self)
        underline_action.triggered.connect(self.underline_text)
        toolbar.addAction(underline_action)

        increase_font_action = QAction(QIcon("increase_font.png"), "Increase Font Size", self)
        increase_font_action.triggered.connect(self.increase_font_size)
        toolbar.addAction(increase_font_action)

        font_combobox = QFontComboBox(self)
        font_combobox.setCurrentFont(QFont("TkDefaultFont"))
        font_combobox.currentFontChanged.connect(self.change_font)
        toolbar.addWidget(font_combobox)

        change_color_action = QAction(QIcon("change_color.png"), "Change Text Color", self)
        change_color_action.triggered.connect(self.change_text_color)
        toolbar.addAction(change_color_action)



    def init_tab_bar(self):
        add_tab_button = QToolButton(self)
        add_tab_button.setText("+")
        add_tab_button.setStyleSheet("QToolButton { font-size: 20px; }")  
        add_tab_button.clicked.connect(self.open_new_tab)
        self.tab_widget.setCornerWidget(add_tab_button)

    def open_file(self):
        file, _ = QFileDialog.getOpenFileName(self, "Open File")
        if file:
            encodings = ["utf-8", "latin-1", "cp1252"]
            content = None
            for encoding in encodings:
                try:
                    with open(file, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            if content is not None:
                current_widget = self.tab_widget.currentWidget()
                current_widget.setPlainText(content)
                self.set_tab_title(current_widget, file) 
            else:
                QMessageBox.warning(self, "Open File", "Unable to open the file.")

    def save_file(self):
        current_widget = self.tab_widget.currentWidget()
        content = current_widget.toPlainText()
        file, _ = QFileDialog.getSaveFileName(self, "Save File")
        if file:
            try:
                with open(file, "w") as f:
                    f.write(content)
                self.set_tab_title(current_widget, file)  # Set the tab title
                QMessageBox.information(self, "Save File", "File saved successfully.")
            except:
                QMessageBox.warning(self, "Save File", "Unable to save the file.")

    def set_tab_title(self, text_widget, file_path):
    # Set the tab title to the file name
        file_name = QFileInfo(file_path).fileName()
        index = self.tab_widget.indexOf(text_widget)
        self.tab_widget.setTabText(index, file_name)
    # Store the file path in the text_widget for later use
        setattr(text_widget, "file_path", file_path)


    def export_as_docx(self):
        file, _ = QFileDialog.getSaveFileName(self, "Export as DOCX", filter="*.docx")
        if file:
            doc = Document()
            current_widget = self.tab_widget.currentWidget()
            content = current_widget.toPlainText()
            paragraph = doc.add_paragraph()
            runs = self.get_runs_with_formatting(content)
            for run_text, font_format in runs:
                run = paragraph.add_run(run_text)
                self.apply_formatting(run, font_format)
            doc.save(file)
            QMessageBox.information(self, "Export as DOCX", "File exported successfully.")

    def export_as_txt(self):
        file, _ = QFileDialog.getSaveFileName(self, "Export as TXT", filter="*.txt")
        if file:
            current_widget = self.tab_widget.currentWidget()
            content = current_widget.toPlainText()
            try:
                with open(file, "w") as f:
                    f.write(content)
                QMessageBox.information(self, "Export as TXT", "File exported successfully.")
            except:
                QMessageBox.warning(self, "Export as TXT", "Unable to export the file.")



    def apply_formatting(self, run, font_format):
        font = run.font
        if font_format["bold"]:
            font.bold = True
        if font_format["italic"]:
            font.italic = True
        if font_format["underline"]:
            font.underline = True
        if font_format["color"]:
            rgb_color = QColor(font_format["color"])
            font.color.rgb = rgb_color.rgb()
        font.size = Pt(15)

    def get_runs_with_formatting(self, text):
        cursor = QTextCursor(self.textarea.document())
        cursor.setPosition(0)
        cursor.movePosition(QTextCursor.End, QTextCursor.KeepAnchor)
        selected_text = cursor.selection().toPlainText()

        runs = []
        start = 0
        for char in selected_text:
            cursor.setPosition(start)
            cursor.movePosition(QTextCursor.NextCharacter, QTextCursor.KeepAnchor)
            char_format = cursor.charFormat()
            runs.append((char, {
                "bold": char_format.font().bold(),
                "italic": char_format.font().italic(),
                "underline": char_format.font().underline(),
                "color": char_format.foreground().color().name()
            }))
            start += 1
        return runs
    
    def closeEvent(self, event):
        current_widget = self.tab_widget.currentWidget()
        if self.is_unsaved_changes(current_widget):
            reply = QMessageBox.question(self, "Unsaved Changes",
                                     "There are unsaved changes. Do you want to save before exiting?",
                                     QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel)
            if reply == QMessageBox.Save:
                self.save_file()
            elif reply == QMessageBox.Cancel:
                event.ignore()
                return
        event.accept()

    def is_unsaved_changes(self, text_widget):
        content = text_widget.toPlainText()
        return content != "" and content != self.get_file_content(text_widget)

    def get_file_content(self, text_widget):
        file_path = getattr(text_widget, "file_path", None)
        if file_path:
            try:
                with open(file_path, "r") as file:
                    content = file.read()
                    return content
            except FileNotFoundError:
                pass
        return ""


 

    def bold_text(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font.setBold(not font.bold())
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def italic_text(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font.setItalic(not font.italic())
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def underline_text(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font.setUnderline(not font.underline())
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def increase_font_size(self):
        current_widget = self.tab_widget.currentWidget()
        cursor = current_widget.textCursor()
        format = cursor.charFormat()
        font = format.font()
        font_size = font.pointSize()
        font_size += 1
        font.setPointSize(font_size)
        format.setFont(font)
        cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def change_font(self, font):
        current_widget = self.tab_widget.currentWidget()
        text_cursor = current_widget.textCursor()
        format = text_cursor.charFormat()

        font_name = font.family()
        format.setFontFamily(font_name)
        text_cursor.mergeCharFormat(format)
        current_widget.setFocus()

    def change_text_color(self):
        color = QColorDialog.getColor(parent=self)
        if color.isValid():
            current_widget = self.tab_widget.currentWidget()
            cursor = current_widget.textCursor()
            format = cursor.charFormat()
            format.setForeground(color)
            cursor.mergeCharFormat(format)
            current_widget.setFocus()



    def open_new_tab(self):
        text_area = QTextEdit()
        text_area.setFont(QFont("TkDefaultFont"))
        text_area.textChanged.connect(self.update_tab_title) 
        self.tab_widget.addTab(text_area, "Untitled")
        self.tab_widget.setCurrentWidget(text_area)

    def close_tab(self, index):
        text_area = self.tab_widget.widget(index)
        if self.is_unsaved_changes(text_area):
            reply = QMessageBox.question(self, "Unsaved Changes",
                                     "There are unsaved changes. Do you want to save before closing the tab?",
                                     QMessageBox.Save | QMessageBox.Discard | QMessageBox.Cancel)
            if reply == QMessageBox.Save:
                self.save_file()
            elif reply == QMessageBox.Cancel:
                return
        self.tab_widget.removeTab(index)

    def update_tab_title(self):
        current_widget = self.tab_widget.currentWidget()
        current_index = self.tab_widget.currentIndex()
        if self.is_unsaved_changes(current_widget):
            file_path = getattr(current_widget, "file_path", None)
            if file_path:
                file_name = os.path.basename(file_path)
                self.tab_widget.setTabText(current_index, file_name +" *")
            else:
                self.tab_widget.setTabText(current_index, "Unsaved *")
        else:
            file_path = getattr(current_widget, "file_path", None)
            if file_path:
                file_name = os.path.basename(file_path)
                self.tab_widget.setTabText(current_index, file_name)
            else:
                self.tab_widget.setTabText(current_index, "Untitled")

    def set_style_options(self):
        style_sheet = """
QTabWidget::pane {
    background-color: #FFFFFF;
}

QTabWidget::tab-bar {
    alignment: left;
    height: auto;
}

QTabBar::tab {
    background-color: #FFFFFF;
    color: #050000;
    border: 1px solid #C0C0C0;
    padding: 10px 0;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    width: 100px;
    height: 10px;
    text-align: center; 
    transition: background-color 0.3s ease, color 0.3s ease;
}

QTabBar::tab:selected {
    background-color: #FFFFFF;
    color: #000000;
    border-bottom-color: #FFFFFF;
}

QTabBar::tab:!selected {
    margin-top: 2px;
    background-color: #ebfcfc;
}

QTabBar::tab:hover {
    background-color: #a5faf8; 
    color: #303030; 
    cursor: pointer; 
}


        """
        self.setStyleSheet(style_sheet)


if __name__ == "__main__":
    app = QApplication([])
    text_editor = TextEditor()
    text_editor.show()
    app.exec_()
