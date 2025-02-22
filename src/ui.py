import codecs
import os
from typing import Optional
import re
import webbrowser
from base64 import b64decode
from urllib.parse import quote

import requests
from PyQt6.QtCore import QSize, QUrl, QFileInfo
from PyQt6.QtGui import (
    QIcon,
    QFont,
    QAction,
    QColor,
    QKeySequence,
    QCloseEvent,
    QKeySequence,
)
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWidgets import (
    QStatusBar,
    QMainWindow,
    QTabWidget,
    QLabel,
    QApplication,
    QToolBar,
    QFontComboBox,
    QToolButton,
    QFileDialog,
    QMessageBox,
    QInputDialog,
)
from docx import Document
from docx.shared import Pt, RGBColor

from widgets.Editor import Editor
from Dialogs.Developers import Developers
from Dialogs.About import About


def get_username_from_about_file():
    with open("user-data/about.txt", "r") as file:
        lines = file.readlines()
        for line in lines:
            if line.startswith("Username:"):
                return line.strip().split(":")[1].strip()
    return None


class TextEditor(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Text Editor")
        self.setGeometry(100, 100, 800, 600)

        self._init_main_layout()
        self._init_status_bar()
        self._init_menu()
        self._init_toolbar()
        self._init_style()

        self.open_empty_tab()

    def _init_main_layout(self):
        self.tab_widget = QTabWidget(self)
        self.setCentralWidget(self.tab_widget)
        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self.close_tab)

        new_tab_button = QToolButton(self)
        new_tab_button.setText("+")
        new_tab_button.setStyleSheet("QToolButton { font-size: 20px; }")
        new_tab_button.clicked.connect(self.open_new_tab)
        self.tab_widget.setCornerWidget(new_tab_button)

    def _init_status_bar(self):
        self.status_bar = QStatusBar()
        self.stats_worl_count = QLabel("Word count: 0", self)
        self.stats_character_count = QLabel("Character count: 0", self)
        self.status_bar.addWidget(self.stats_worl_count)
        self.status_bar.addWidget(self.stats_character_count)
        self.line_number_label = QLabel()
        self.status_bar.addPermanentWidget(self.line_number_label)
        self.setStatusBar(self.status_bar)

    def _init_toolbar(self):
        toolbar = QToolBar("Format Toolbar", self)
        self.addToolBar(toolbar)

        toolbar.setIconSize(QSize(20, 20))

        bold_action = QAction(QIcon("images/bold.png"), "Bold", self)
        bold_action.triggered.connect(self.toggle_editor_bold)
        toolbar.addAction(bold_action)

        italic_action = QAction(QIcon("images/italic.png"), "Italic", self)
        italic_action.triggered.connect(self.toggle_editor_italic)
        toolbar.addAction(italic_action)

        underline_action = QAction(QIcon("images/underline.png"), "Underline", self)
        underline_action.triggered.connect(self.toggle_editor_underline)
        toolbar.addAction(underline_action)

        increase_font_action = QAction(
            QIcon("images/increase_font.png"), "Increase Font Size", self
        )
        increase_font_action.triggered.connect(self.increase_editor_font_size)
        toolbar.addAction(increase_font_action)

        decrease_font_action = QAction(
            QIcon("images/decrease_font.png"), "Decrease Font Size", self
        )
        decrease_font_action.triggered.connect(self.decrease_editor_font_size)
        toolbar.addAction(decrease_font_action)

        font_combobox = QFontComboBox(self)
        font_combobox.setCurrentFont(QFont("TkDefaultFont"))
        font_combobox.currentFontChanged.connect(self.change_editor_font)
        toolbar.addWidget(font_combobox)

        change_color_action = QAction(
            QIcon("images/change_color.png"), "Change Text Color", self
        )
        change_color_action.triggered.connect(self.ask_editor_text_color)
        toolbar.addAction(change_color_action)

        set_text_background_color = QAction(
            QIcon("images/change_bg_color.png"), "Change Background Color", self
        )
        set_text_background_color.triggered.connect(self.ask_editor_text_bg_color)
        toolbar.addAction(set_text_background_color)

    def _init_menu(self):
        menubar = self.menuBar()
        assert menubar is not None

        info_menu = menubar.addMenu("Infos")
        assert info_menu is not None
        developer_action = QAction("Developer", self)
        developer_action.triggered.connect(self.show_developers)
        info_menu.addAction(developer_action)

        about_action = QAction("About", self)
        about_action.triggered.connect(self.show_about)
        info_menu.addAction(about_action)

        file_menu = menubar.addMenu("File")
        assert file_menu is not None
        open_action = QAction("Open", self)
        open_action.triggered.connect(self.open_file)
        file_menu.addAction(open_action)

        new_tab_action = QAction("New Tab", self)
        new_tab_action.setShortcut(QKeySequence("Ctrl+T"))
        new_tab_action.triggered.connect(self.open_new_tab)
        file_menu.addAction(new_tab_action)

        search_action = QAction("Search Word", self)
        search_action.setShortcut(QKeySequence("Ctrl+F"))
        search_action.triggered.connect(self.open_editor_find)
        file_menu.addAction(search_action)

        exit_action = QAction("Exit", self)
        exit_action.triggered.connect(QApplication.quit)
        file_menu.addAction(exit_action)

        save_menu = menubar.addMenu("Save")
        assert save_menu is not None
        save_action = QAction("Save", self)
        save_action.setShortcut(QKeySequence("Ctrl+S"))
        save_action.triggered.connect(self.save_file)
        save_menu.addAction(save_action)

        export_docx_action = QAction("Export as DOCX", self)
        export_docx_action.triggered.connect(self.export_as_docx)
        save_menu.addAction(export_docx_action)

        export_txt_action = QAction("Export as TXT", self)
        export_txt_action.triggered.connect(self.export_as_txt)
        save_menu.addAction(export_txt_action)

        print_action = QAction("Print", self)
        print_action.setShortcut(QKeySequence("Ctrl+P"))
        print_action.triggered.connect(self.print_editor)
        save_menu.addAction(print_action)

        edit_menu = menubar.addMenu("Edit")
        assert edit_menu is not None
        undo_action = QAction("Undo", self)
        undo_action.setShortcut(QKeySequence.StandardKey.Undo)
        undo_action.triggered.connect(self.undo)
        edit_menu.addAction(undo_action)

        redo_action = QAction("Redo", self)
        redo_action.setShortcut(QKeySequence.StandardKey.Redo)
        redo_action.triggered.connect(self.redo)
        edit_menu.addAction(redo_action)

        projects_menu = menubar.addMenu("Projects")
        assert projects_menu is not None
        login_action = QAction("Login", self)
        login_action.triggered.connect(self.start_webserver)
        projects_menu.addAction(login_action)

        projects = self.load_projects()
        if projects:
            projects_menu.addSeparator()

        for project in projects:
            project_action = QAction(project, self)
            project_action.triggered.connect(lambda _, p=project: self.open_project(p))
            projects_menu.addAction(project_action)

    def update_status_bar(self):
        current_widget = self.tab_widget.currentWidget()
        if isinstance(current_widget, Editor):
            content = current_widget.toPlainText()

            word_count = len(re.findall(r"\b\w+\'?\w*\b", content))
            char_count = len(content)

            cursor = current_widget.textCursor()
            line_number = cursor.blockNumber() + 1

            self.stats_worl_count.setText(f"Word count: {word_count}")
            self.stats_character_count.setText(f"Character count: {char_count}")
            self.line_number_label.setText(f"Line: {line_number}")
        else:
            self.stats_worl_count.setText("")
            self.stats_character_count.setText("")
            self.line_number_label.setText("")

    def show_developers(self):
        developers_dialog = Developers(self)
        developers_dialog.exec()

    def show_about(self):
        about_dialog = About(self)
        about_dialog.exec()

    def open_project(self, project):
        username = get_username_from_about_file()
        if username:
            repo_url = f"https://api.github.com/repos/{username}/{project}/contents"
            try:
                response = requests.get(repo_url)
                response.raise_for_status()

                files = [file_info["name"] for file_info in response.json()]

                if files:
                    selected_file, ok = QInputDialog.getItem(
                        self, "Select File", "Select a file to open:", files, 0, False
                    )
                    if ok and selected_file:
                        file_content_url = f"https://api.github.com/repos/{username}/{project}/contents/{quote(selected_file)}"
                        content_response = requests.get(file_content_url)
                        content_response.raise_for_status()

                        content = b64decode(content_response.json()["content"]).decode(
                            "utf-8"
                        )

                        current_widget = self.tab_widget.currentWidget()
                        assert isinstance(current_widget, Editor)
                        current_widget.setPlainText(content)
                        self.set_tab_title(current_widget, selected_file)
                else:
                    QMessageBox.warning(
                        self, "No Files", f"There are no files in {project}."
                    )
            except requests.RequestException as e:
                QMessageBox.warning(
                    self, "Error", f"Error fetching project files: {str(e)}"
                )
        else:
            pass

    def load_projects(self):
        projects = []
        with open("user-data/projects.txt", "r") as file:
            for line in file:
                project = line.strip()
                if project:
                    projects.append(project)
        return projects

    def start_webserver(self):
        url = "http://127.0.0.1:5000/login"
        webbrowser.open(url)

    def undo(self):
        current_widget = self.tab_widget.currentWidget()
        assert isinstance(current_widget, Editor)
        current_widget.undo()

    def redo(self):
        current_widget = self.tab_widget.currentWidget()
        assert isinstance(current_widget, Editor)
        current_widget.redo()

    def open_file(self):
        file, _ = QFileDialog.getOpenFileName(self, "Open File")
        if file:
            encodings = ["utf-8", "latin-1", "cp1252"]
            content = None
            for encoding in encodings:
                try:
                    with open(file, "r", encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            if content is not None:
                current_widget = self.tab_widget.currentWidget()
                assert isinstance(current_widget, Editor)
                current_widget.setPlainText(content)
                self.set_tab_title(current_widget, file)
            else:
                QMessageBox.warning(self, "Open File", "Unable to open the file.")

    def save_file(self):
        options = ["Save locally", "Save on GitHub"]
        choice, ok = QInputDialog.getItem(
            self, "Save Option", "Choose a save option:", options, 0, False
        )

        if ok:
            current_widget = self.tab_widget.currentWidget()
            assert isinstance(current_widget, Editor)
            content = current_widget.toPlainText()

            if choice == "Save locally":
                self.save_locally(content)
            elif choice == "Save on GitHub":
                if os.path.exists("user-data/upload_data.txt"):
                    github_username, access_token = self.read_upload_data()
                    repo_name = self.get_user_repository(github_username, access_token)
                    if repo_name:
                        self.save_to_github(
                            content, github_username, access_token, repo_name
                        )
                    else:
                        QMessageBox.warning(
                            self,
                            "GitHub Save",
                            "No repositories found for the given GitHub username.",
                        )
                else:
                    self.ask_github_credentials_and_save(content)

    def save_locally(self, content):
        file_dialog = QFileDialog()
        file_name, _ = file_dialog.getSaveFileName(
            self, "Save File Locally", "", "All Files (*)"
        )

        if file_name:
            with open(file_name, "w") as f:
                f.write(content)
            QMessageBox.information(
                self, "Save Successful", "File saved locally successfully."
            )

    def ask_github_credentials_and_save(self, content):
        github_username, ok1 = QInputDialog.getText(
            self, "GitHub Credentials", "Enter your GitHub username:"
        )
        access_token, ok2 = QInputDialog.getText(
            self, "GitHub Credentials", "Enter your GitHub personal access token:"
        )

        if ok1 and ok2:
            repositories = self.get_user_repositories(github_username, access_token)

            if repositories:
                repository_name, ok3 = QInputDialog.getItem(
                    self,
                    "Select Repository",
                    "Choose a GitHub repository:",
                    repositories,
                    0,
                    False,
                )

                if ok3:
                    with open("user-data/upload_data.txt", "w") as file:
                        file.write(
                            f"{github_username}\n{access_token}\n{repository_name}"
                        )

                    self.save_to_github(
                        content, github_username, access_token, repository_name
                    )
                else:
                    QMessageBox.warning(
                        self,
                        "GitHub Save",
                        "No GitHub repository selected. Unable to save on GitHub.",
                    )
            else:
                QMessageBox.warning(
                    self,
                    "GitHub Save",
                    "No repositories found for the given GitHub username.",
                )
        else:
            QMessageBox.warning(
                self,
                "GitHub Save",
                "GitHub credentials not provided. Unable to save on GitHub.",
            )

    def save_to_github(self, content, github_username, access_token, repo_name):
        custom_github_filename, ok = QInputDialog.getText(
            self,
            "GitHub File Name",
            "Enter the desired file name for GitHub (with extension):",
        )

        if ok and custom_github_filename:
            api_url = f"https://api.github.com/repos/{github_username}/{repo_name}/contents/{custom_github_filename}"
            headers = {"Authorization": f"token {access_token}"}

            data = {
                "message": "Upload file via schBenedikt's Text Editor",
                "content": codecs.encode(content.encode("utf-8"), "base64").decode(
                    "utf-8"
                ),
                "sha": self.get_sha_from_github(
                    custom_github_filename, github_username, access_token, repo_name
                ),
            }

            response = requests.put(api_url, headers=headers, json=data)

            if response.status_code == 200:
                print(
                    f"File '{custom_github_filename}' uploaded to GitHub successfully."
                )
            else:
                print(
                    f"Unable to upload file to GitHub. Status code: {response.status_code}, Message: {response.text}"
                )

    def get_user_repositories(self, github_username, access_token):
        api_url = f"https://api.github.com/users/{github_username}/repos"
        headers = {"Authorization": f"token {access_token}"}

        try:
            response = requests.get(api_url, headers=headers)
            repositories = [repo["name"] for repo in response.json()]
            return repositories
        except requests.RequestException as e:
            print(f"Error getting GitHub repositories: {e}")
            return None

    def get_user_repository(self, github_username, access_token):
        repositories = self.get_user_repositories(github_username, access_token)
        if repositories:
            repository_name, ok = QInputDialog.getItem(
                self,
                "Select Repository",
                "Choose a GitHub repository:",
                repositories,
                0,
                False,
            )
            if ok:
                return repository_name
        return None

    def read_upload_data(self):
        with open("user-data/upload_data.txt", "r") as file:
            lines = file.readlines()
            github_username = lines[0].strip()
            access_token = lines[1].strip()
            return github_username, access_token

    def upload_to_github(
        self, content, github_filename, github_username, access_token, repo_name
    ):
        github_filename = os.path.basename(github_filename)

        api_url = f"https://api.github.com/repos/{github_username}/{repo_name}/contents/{github_filename}"
        headers = {"Authorization": f"token {access_token}"}

        data = {
            "message": "Update file via script",
            "content": codecs.encode(content.encode("utf-8"), "base64").decode("utf-8"),
            "sha": self.get_sha_from_github(
                github_filename, github_username, access_token, repo_name
            ),
        }

        response = requests.put(api_url, headers=headers, json=data)

        if response.status_code == 200:
            print(f"File '{github_filename}' uploaded to GitHub successfully.")
        else:
            print(
                f"Unable to upload file to GitHub. Status code: {response.status_code}, Message: {response.text}"
            )

    def get_sha_from_github(
        self, github_filename, github_username, access_token, repo_name
    ):
        api_url = f"https://api.github.com/repos/{github_username}/{repo_name}/contents/{github_filename}"
        headers = {"Authorization": f"token {access_token}"}

        response = requests.get(api_url, headers=headers)

        if response.status_code == 200:
            sha = response.json().get("sha")
            return sha
        else:
            print(
                f"Unable to get SHA from GitHub. Status code: {response.status_code}, Message: {response.text}"
            )
            return None

    def load_github_credentials(self):
        try:
            with open("user-data/upload_data.txt", "r") as file:
                lines = file.readlines()
                github_username = lines[0].strip()
                access_token = lines[1].strip()
                repo_name = lines[2].strip()
                return github_username, access_token, repo_name
        except FileNotFoundError:
            print("upload_data.txt not found.")
            return None, None, None

    def set_tab_title(self, text_widget, file_path):
        file_name = QFileInfo(file_path).fileName()
        index = self.tab_widget.indexOf(text_widget)
        self.tab_widget.setTabText(index, file_name)
        setattr(text_widget, "file_path", file_path)

    def export_as_docx(self):
        file, _ = QFileDialog.getSaveFileName(self, "Export as DOCX", filter="*.docx")
        if file:
            doc = Document()
            editor_widget = self.get_current_editor()
            assert editor_widget is not None
            paragraph = doc.add_paragraph()
            runs = editor_widget.get_runs_with_formatting()
            for run_text, font_format in runs:
                run = paragraph.add_run(run_text)
                self.apply_formatting(run, font_format)
            doc.save(file)
            QMessageBox.information(
                self, "Export as DOCX", "File exported successfully."
            )

    def export_as_txt(self):
        file, _ = QFileDialog.getSaveFileName(self, "Export as TXT", filter="*.txt")
        if file:
            current_widget = self.tab_widget.currentWidget()
            assert isinstance(current_widget, Editor)

            content = current_widget.toPlainText()
            try:
                with open(file, "w") as f:
                    f.write(content)
                QMessageBox.information(
                    self, "Export as TXT", "File exported successfully."
                )
            except:
                QMessageBox.warning(self, "Export as TXT", "Unable to export the file.")

    def apply_formatting(self, run, font_format):
        font = run.font
        if font_format["bold"]:
            font.bold = True
        if font_format["italic"]:
            font.italic = True
        if font_format["underline"]:
            font.underline = True
        if font_format["color"]:
            rgb_color = QColor(font_format["color"])
            font.color.rgb = RGBColor(
                rgb_color.red(), rgb_color.green(), rgb_color.blue()
            )
        font.size = Pt(15)

    def is_unsaved_changes(self, text_widget):
        if isinstance(text_widget, Editor):
            content = text_widget.toPlainText()
            return content != "" and content != self.get_file_content(text_widget)
        return False

    def get_file_content(self, text_widget):
        file_path = getattr(text_widget, "file_path", None)
        if file_path:
            try:
                with open(file_path, "r") as file:
                    content = file.read()
                    return content
            except FileNotFoundError:
                pass
        return ""

    def open_new_tab(self):
        options = ["New File", "Open File", "Chat"]
        selected_option, ok = QInputDialog.getItem(
            self,
            "New File, Open File, or Chat?",
            "Choose an option:",
            options,
            0,
            False,
        )

        if ok:
            if selected_option == "New File":
                dialog = QFileDialog(self)
                dialog.setFileMode(QFileDialog.FileMode.AnyFile)

                options = QFileDialog.Option.DontUseNativeDialog
                fileName, _ = dialog.getSaveFileName(
                    self,
                    "New File",
                    "",
                    "All Files (*);;Text Files (*.txt);;Python Files (*.py)",
                    options=options,
                )

                if fileName:
                    self.open_text_file_in_tab(fileName)
                else:
                    self.open_empty_tab()
            elif selected_option == "Open File":
                dialog = QFileDialog(self)
                dialog.setFileMode(QFileDialog.FileMode.ExistingFile)

                options = QFileDialog.Option.DontUseNativeDialog
                fileName, _ = dialog.getOpenFileName(
                    self,
                    "Open File",
                    "",
                    "All Files (*);;Text Files (*.txt);;Python Files (*.py)",
                    options=options,
                )

                if fileName:
                    self.open_text_file_in_tab(fileName)
            elif selected_option == "Chat":
                self.open_chat_tab()

    def toggle_editor_bold(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.bold_text()

    def toggle_editor_italic(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.italic_text()

    def toggle_editor_underline(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.underline_text()

    def change_editor_font(self, font: QFont):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.change_font(font)

    def ask_editor_text_color(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.ask_for_text_color()

    def ask_editor_text_bg_color(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.ask_for_text_bg_color()

    def decrease_editor_font_size(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.decrease_font_size()

    def increase_editor_font_size(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.increase_font_size()

    def open_editor_find(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.open_search()

    def print_editor(self):
        current_editor = self.get_current_editor()
        assert current_editor is not None
        current_editor.print_document()

    def open_chat_tab(self):
        chat_view = QWebEngineView()
        chat_view.setUrl(QUrl("https://platform.openai.com/"))
        self.tab_widget.addTab(chat_view, "Chat")

    def open_text_file_in_tab(self, file_path):
        with open(file_path, "r") as file:
            content = file.read()

        text_area = Editor()
        text_area.setFont(QFont("TkDefaultFont"))
        text_area.textChanged.connect(self.update_tab_title)
        text_area.setPlainText(content)
        self.tab_widget.addTab(text_area, os.path.basename(file_path))
        self.tab_widget.setCurrentWidget(text_area)

    def open_empty_tab(self):
        text_area = Editor()
        text_area.setFont(QFont("TkDefaultFont"))
        text_area.textChanged.connect(self.update_tab_title)
        text_area.textChanged.connect(self.update_status_bar)
        self.tab_widget.addTab(text_area, "Untitled")
        self.tab_widget.setCurrentWidget(text_area)

    def open_new_empty_tab(self):
        text_area = Editor()
        text_area.setFont(QFont("TkDefaultFont"))
        text_area.textChanged.connect(self.update_tab_title)
        text_area.textChanged.connect(self.update_status_bar)
        self.tab_widget.addTab(text_area, "Untitled")
        self.tab_widget.setCurrentWidget(text_area)

    def close_tab(self, index):
        text_area = self.tab_widget.widget(index)

        if self.is_unsaved_changes(text_area):
            reply = QMessageBox.question(
                self,
                "Unsaved Changes",
                "There are unsaved changes. Do you want to save before closing the tab?",
                QMessageBox.StandardButton.Save
                | QMessageBox.StandardButton.Discard
                | QMessageBox.StandardButton.Cancel,
            )

            if reply == QMessageBox.StandardButton.Save:
                self.save_file()
            elif reply == QMessageBox.StandardButton.Cancel:
                return

        self.tab_widget.removeTab(index)

    def closeEvent(self, a0: "QCloseEvent | None") -> None:
        if a0 is None:
            return

        current_widget = self.tab_widget.currentWidget()
        if self.is_unsaved_changes(current_widget):
            reply = QMessageBox.question(
                self,
                "Unsaved Changes",
                "There are unsaved changes. Do you want to save before exiting?",
                QMessageBox.StandardButton.Save
                | QMessageBox.StandardButton.Discard
                | QMessageBox.StandardButton.Cancel,
            )
            if reply == QMessageBox.StandardButton.Save:
                self.save_file()
            elif reply == QMessageBox.StandardButton.Cancel:
                a0.ignore()
                return
        a0.accept()
        if self.tab_widget.count() == 0:
            self.close()

    def update_tab_title(self):
        current_widget = self.tab_widget.currentWidget()
        current_index: int = self.tab_widget.currentIndex()
        if self.is_unsaved_changes(current_widget):
            file_path = getattr(current_widget, "file_path", None)
            if file_path:
                file_name = os.path.basename(file_path)
                self.tab_widget.setTabText(current_index, file_name + " *")
            else:
                self.tab_widget.setTabText(current_index, "Unsaved *")
        else:
            file_path = getattr(current_widget, "file_path", None)
            if file_path:
                file_name = os.path.basename(file_path)
                self.tab_widget.setTabText(current_index, file_name)
            else:
                self.tab_widget.setTabText(current_index, "Untitled")

    def get_current_editor(self) -> Optional[Editor]:
        current_widget = self.tab_widget.currentWidget()
        return current_widget if isinstance(current_widget, Editor) else None

    def _init_style(self):
        style_sheet = """
QTabWidget::pane {
    background-color: #FFFFFF;
}

QTabWidget::tab-bar {
    alignment: left;
    height: auto;
}

QTabBar::tab {
    background-color: #FFFFFF;
    color: #050000;
    border: 1px solid #C0C0C0;
    padding: 10px 0;
    border-top-left-radius: 4px;
    border-top-right-radius: 4px;
    width: 100px;
    height: 10px;
    text-align: center; 
}

QTabBar::tab:selected {
    background-color: #FFFFFF;
    color: #000000;
    border-bottom-color: #FFFFFF;
}

QTabBar::tab:!selected {
    margin-top: 2px;
    background-color: #ebfcfc;
}

QTabBar::tab:hover {
    background-color: #a5faf8; 
    color: #303030; 
}
        """
        self.setStyleSheet(style_sheet)
